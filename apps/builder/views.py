from io import BytesIO
from urllib.parse import quote

from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.utils.http import url_has_allowed_host_and_scheme
from django.views.decorators.http import require_POST

from apps.accounts.decorators import analyst_required
from apps.knowledge.models import Norm, CourtCase


def _get_basket(request):
    return request.session.setdefault('basket', {'norms': [], 'cases': []})


def _save_basket(request, basket):
    request.session['basket'] = basket
    request.session.modified = True


@analyst_required
def builder_home(request):
    basket = _get_basket(request)
    norms = Norm.objects.filter(pk__in=basket['norms']).select_related('branch')
    cases = CourtCase.objects.filter(pk__in=basket['cases']).select_related('branch')
    return render(request, 'builder/builder.html', {'norms': norms, 'cases': cases})


@analyst_required
@require_POST
def basket_add(request, model_type, pk):
    basket = _get_basket(request)
    if model_type == 'norm':
        get_object_or_404(Norm, pk=pk)
        if pk not in basket['norms']:
            basket['norms'].append(pk)
            _save_basket(request, basket)
    elif model_type == 'case':
        get_object_or_404(CourtCase, pk=pk)
        if pk not in basket['cases']:
            basket['cases'].append(pk)
            _save_basket(request, basket)

    next_url = request.POST.get('next') or '/'
    if not url_has_allowed_host_and_scheme(next_url, allowed_hosts={request.get_host()}):
        next_url = '/'
    return redirect(next_url)


@analyst_required
@require_POST
def basket_remove(request, model_type, pk):
    basket = _get_basket(request)
    if model_type == 'norm':
        basket['norms'] = [i for i in basket['norms'] if i != pk]
    elif model_type == 'case':
        basket['cases'] = [i for i in basket['cases'] if i != pk]
    _save_basket(request, basket)

    if request.headers.get('HX-Request'):
        updated = _get_basket(request)
        norms = Norm.objects.filter(pk__in=updated['norms']).select_related('branch')
        cases = CourtCase.objects.filter(pk__in=updated['cases']).select_related('branch')
        return render(request, 'builder/_basket.html', {'norms': norms, 'cases': cases})
    return redirect('builder:home')


@analyst_required
@require_POST
def basket_clear(request):
    request.session['basket'] = {'norms': [], 'cases': []}
    request.session.modified = True

    if request.headers.get('HX-Request'):
        return render(request, 'builder/_basket.html', {'norms': [], 'cases': []})
    return redirect('builder:home')


@analyst_required
@require_POST
def export_docx(request):
    basket = _get_basket(request)
    norms = list(Norm.objects.filter(pk__in=basket['norms']).select_related('branch'))
    cases = list(CourtCase.objects.filter(pk__in=basket['cases']).select_related('branch'))

    title = request.POST.get('title', '').strip() or 'Правовое заключение'
    question = request.POST.get('question', '').strip()
    conclusion = request.POST.get('conclusion', '').strip()

    doc = _build_docx(title, question, norms, cases, conclusion, request.user)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = ''.join(c for c in title if c.isalnum() or c in ' _-')[:50].strip() or 'zaklyuchenie'
    encoded_name = quote(f'{safe_name}.docx', safe='')
    response = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f"attachment; filename=\"zaklyuchenie.docx\"; filename*=UTF-8''{encoded_name}"
    return response


# ─── docx builder ────────────────────────────────────────────────────────────

def _build_docx(title, question, norms, cases, conclusion, user):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.utils import timezone

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Title
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Meta
    date_str = timezone.now().strftime('%d.%m.%Y')
    author_name = user.get_full_name() or user.username
    _para(doc, f'Дата составления: {date_str}')
    _para(doc, f'Составитель: {author_name}')

    doc.add_paragraph()

    # Question
    if question:
        _heading(doc, 'Правовой вопрос')
        _para(doc, question)
        doc.add_paragraph()

    # Norms
    if norms:
        _heading(doc, 'Применимые нормы права')
        for norm in norms:
            _norm_block(doc, norm)
        doc.add_paragraph()

    # Cases
    if cases:
        _heading(doc, 'Судебная практика')
        for case in cases:
            _case_block(doc, case)
        doc.add_paragraph()

    # Conclusion
    _heading(doc, 'Заключение')
    _para(doc, conclusion or 'Текст заключения не заполнен.')

    return doc


def _heading(doc, text):
    from docx.shared import Pt
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(13)


def _para(doc, text):
    from docx.shared import Pt, Cm
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.runs[0].font.size = Pt(12)
    return p


def _norm_block(doc, norm):
    from docx.shared import Pt, Cm
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(norm.article or norm.title)
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f' ({norm.get_norm_type_display()})').font.size = Pt(11)

    if norm.text:
        excerpt = norm.text[:500] + ('…' if len(norm.text) > 500 else '')
        tp = doc.add_paragraph(excerpt)
        tp.paragraph_format.left_indent = Cm(1)
        tp.runs[0].font.size = Pt(11)


def _case_block(doc, case):
    from docx.shared import Pt, Cm
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(case.case_number)
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f' — {case.court}, {case.decision_date.strftime("%d.%m.%Y")}').font.size = Pt(11)

    if case.thesis:
        tp = doc.add_paragraph(case.thesis)
        tp.paragraph_format.left_indent = Cm(1)
        tp.runs[0].font.size = Pt(11)
