"""Скрипт для генерации .docx-шаблонов с docxtpl-тегами. Запускать вручную при необходимости."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent


def _margins(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(3)
    s.right_margin = Cm(1.5)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.size = Pt(12 if level == 1 else 11)
    return p


def _center(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def _подписи(doc, left_label, left_body, right_label, right_body):
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = left_label
    t.rows[0].cells[1].text = right_label
    r = t.add_row().cells
    r[0].text = left_body
    r[1].text = right_body


# ──────────────────────────────────────────────────────────────────────────────
# 1. Договор оказания услуг (уже существует, не пересоздаём)
# ──────────────────────────────────────────────────────────────────────────────

def create_dogovor_uslug():
    doc = Document()
    _margins(doc)

    _center(doc, 'ДОГОВОР ОКАЗАНИЯ УСЛУГ').runs[0].bold = True
    _center(doc, '№ {{ number }}  от  {{ data }} г.')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('{{ contragent_name }}')
    p.add_run(
        ', именуемое в дальнейшем «Заказчик», в лице директора {{ contragent_director }}, '
        'действующего на основании Устава, с одной стороны, и '
    )
    p.add_run('{{ executor_name }}')
    p.add_run(
        ', именуемое в дальнейшем «Исполнитель», в лице {{ lawyer_name }}, '
        'с другой стороны, заключили настоящий договор о нижеследующем:'
    )

    _heading(doc, '1. ПРЕДМЕТ ДОГОВОРА', 2)
    doc.add_paragraph('1.1. Исполнитель обязуется оказать Заказчику следующие услуги: {{ services_description }}.')
    doc.add_paragraph('1.2. Услуги оказываются по адресу: {{ contragent_address }}.')

    _heading(doc, '2. СТОИМОСТЬ И ПОРЯДОК РАСЧЁТОВ', 2)
    doc.add_paragraph(
        '2.1. Стоимость услуг составляет {{ summa }} руб. ({{ summa_words }}).'
    )
    doc.add_paragraph(
        '2.2. Оплата производится в течение {{ payment_days }} рабочих дней после подписания акта.'
    )

    _heading(doc, '3. СРОК ОКАЗАНИЯ УСЛУГ', 2)
    doc.add_paragraph('3.1. Срок оказания услуг: с {{ date_start }} по {{ date_end }}.')

    _heading(doc, '4. РЕКВИЗИТЫ СТОРОН', 2)
    _подписи(doc,
        'ЗАКАЗЧИК',
        '{{ contragent_name }}\nИНН {{ contragent_inn }}\nКПП {{ contragent_kpp }}\n'
        'Адрес: {{ contragent_address }}\n\nПодпись: _______________\n{{ contragent_director }}',
        'ИСПОЛНИТЕЛЬ',
        '{{ executor_name }}\n\nПодпись: _______________\n{{ lawyer_name }}',
    )

    path = OUT / 'dogovor_uslug.docx'
    doc.save(path)
    print(f'✓ {path.name}')
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 2. Договор поставки
# ──────────────────────────────────────────────────────────────────────────────

def create_dogovor_postavki():
    doc = Document()
    _margins(doc)

    _center(doc, 'ДОГОВОР ПОСТАВКИ').runs[0].bold = True
    _center(doc, '№ {{ number }}  от  {{ data }} г.')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('{{ contragent_name }}')
    p.add_run(
        ', именуемое в дальнейшем «Покупатель», в лице директора {{ contragent_director }}, '
        'и '
    )
    p.add_run('{{ executor_name }}')
    p.add_run(
        ', именуемое в дальнейшем «Поставщик», в лице {{ lawyer_name }}, '
        'заключили настоящий договор о нижеследующем:'
    )

    _heading(doc, '1. ПРЕДМЕТ ДОГОВОРА', 2)
    doc.add_paragraph(
        '1.1. Поставщик обязуется поставить Покупателю товар: {{ tovar_naim }}, '
        'в количестве {{ tovar_kol }} ед.'
    )
    doc.add_paragraph('1.2. Адрес доставки: {{ delivery_address }}.')
    doc.add_paragraph('1.3. Срок поставки: {{ delivery_date }}.')

    _heading(doc, '2. ЦЕНА И ПОРЯДОК ОПЛАТЫ', 2)
    doc.add_paragraph(
        '2.1. Цена единицы товара составляет {{ tovar_price }} руб.'
    )
    doc.add_paragraph(
        '2.2. Общая стоимость договора: {{ tovar_summa }} руб. ({{ summa_words }}).'
    )
    doc.add_paragraph(
        '2.3. Оплата производится в течение {{ payment_days }} рабочих дней с момента поставки.'
    )

    _heading(doc, '3. ОТВЕТСТВЕННОСТЬ СТОРОН', 2)
    doc.add_paragraph(
        '3.1. За просрочку поставки Поставщик уплачивает пеню в размере 0,1 % от суммы '
        'договора за каждый день просрочки.'
    )

    _heading(doc, '4. РЕКВИЗИТЫ СТОРОН', 2)
    _подписи(doc,
        'ПОКУПАТЕЛЬ',
        '{{ contragent_name }}\nИНН {{ contragent_inn }}\nКПП {{ contragent_kpp }}\n'
        'Адрес: {{ contragent_address }}\n\nПодпись: _______________\n{{ contragent_director }}',
        'ПОСТАВЩИК',
        '{{ executor_name }}\n\nПодпись: _______________\n{{ lawyer_name }}',
    )

    path = OUT / 'dogovor_postavki.docx'
    doc.save(path)
    print(f'✓ {path.name}')
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 3. Претензия
# ──────────────────────────────────────────────────────────────────────────────

def create_pretenziya():
    doc = Document()
    _margins(doc)

    # Шапка
    p_right = doc.add_paragraph()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run(
        'Генеральному директору\n{{ contragent_name }}\n{{ contragent_address }}\n\n'
        'от {{ sender_name }}\n{{ sender_address }}'
    )

    doc.add_paragraph()
    _center(doc, 'ПРЕТЕНЗИЯ').runs[0].bold = True
    _center(doc, '№ {{ number }}  от  {{ data }} г.')
    doc.add_paragraph()

    doc.add_paragraph(
        'Между нашей организацией и {{ contragent_name }} (ИНН {{ contragent_inn }}) '
        'заключён договор № {{ dogovor_number }} от {{ dogovor_date }} г.'
    )
    doc.add_paragraph(
        'В соответствии с условиями указанного договора {{ contragent_name }} '
        'приняло на себя обязательство {{ obyazatelstvo }}. '
        'Однако данное обязательство не было исполнено надлежащим образом: {{ narushenie }}.'
    )
    doc.add_paragraph(
        'В связи с изложенным, на основании ст. 309, 310, 393 ГК РФ, требуем: {{ trebovanie }}.'
    )
    doc.add_paragraph(
        'Сумма претензии составляет {{ pretenziya_summa }} руб. ({{ summa_words }}).'
    )
    doc.add_paragraph(
        'Просим рассмотреть настоящую претензию и дать ответ в течение {{ srok_otveta }} '
        'рабочих дней с момента получения.'
    )
    doc.add_paragraph(
        'В случае неудовлетворения претензии в добровольном порядке '
        'оставляем за собой право обратиться в суд.'
    )

    doc.add_paragraph()
    doc.add_paragraph(
        '{{ lawyer_name }}  _______________ / {{ data }} г.'
    )

    path = OUT / 'pretenziya.docx'
    doc.save(path)
    print(f'✓ {path.name}')
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 4. Доверенность
# ──────────────────────────────────────────────────────────────────────────────

def create_doverennost():
    doc = Document()
    _margins(doc)

    _center(doc, 'ДОВЕРЕННОСТЬ').runs[0].bold = True
    _center(doc, '№ {{ number }}').runs[0].bold = False
    doc.add_paragraph()

    doc.add_paragraph(
        'г. {{ gorod }},  {{ data }} г.'
    )
    doc.add_paragraph(
        '{{ org_name }}, ИНН {{ org_inn }}, расположенное по адресу: {{ org_address }}, '
        'в лице директора {{ director_name }}, действующего на основании Устава, '
        'настоящей доверенностью уполномочивает:'
    )
    doc.add_paragraph(
        '{{ poverenny_name }}, паспорт {{ poverenny_pasport }}, '
        'зарегистрированного по адресу: {{ poverenny_address }},'
    )
    doc.add_paragraph(
        'представлять интересы {{ org_name }} по следующим вопросам: {{ polnomochiya }}.'
    )
    doc.add_paragraph(
        'Доверенность выдана сроком до {{ srok_do }}.'
    )
    doc.add_paragraph('Подпись доверенного лица: _______________')
    doc.add_paragraph()
    doc.add_paragraph(
        '{{ director_name }}  _______________'
    )

    path = OUT / 'doverennost.docx'
    doc.save(path)
    print(f'✓ {path.name}')
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 5. Исковое заявление
# ──────────────────────────────────────────────────────────────────────────────

def create_iskovoe():
    doc = Document()
    _margins(doc)

    p_right = doc.add_paragraph()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run(
        'В {{ sud_name }}\n\n'
        'Истец: {{ istec_name }}\nАдрес: {{ istec_address }}\nИНН: {{ istec_inn }}\n\n'
        'Ответчик: {{ otvetchik_name }}\nАдрес: {{ otvetchik_address }}\nИНН: {{ otvetchik_inn }}\n\n'
        'Цена иска: {{ summa_iska }} руб.'
    )

    doc.add_paragraph()
    _center(doc, 'ИСКОВОЕ ЗАЯВЛЕНИЕ').runs[0].bold = True
    _center(doc, '{{ data }} г.')
    doc.add_paragraph()

    doc.add_paragraph(
        'Между {{ istec_name }} и {{ otvetchik_name }} заключён договор № {{ dogovor_number }} '
        'от {{ dogovor_date }} г. (далее — Договор).'
    )
    doc.add_paragraph(
        'В нарушение условий Договора ответчик {{ narushenie }}.'
    )
    doc.add_paragraph(
        'Истцом соблюдён досудебный порядок урегулирования спора: '
        'претензия направлена {{ pretenziya_date }} г., ответ не получен.'
    )

    _heading(doc, 'ИСКОВЫЕ ТРЕБОВАНИЯ', 2)
    doc.add_paragraph('На основании изложенного, руководствуясь ст. 4, 125, 126 АПК РФ, прошу:')
    doc.add_paragraph('{{ iskovye_trebovaniya }}')
    doc.add_paragraph(
        'Взыскать с ответчика {{ summa_iska }} руб. ({{ summa_words }}).'
    )

    _heading(doc, 'ПРИЛОЖЕНИЯ', 2)
    doc.add_paragraph('1. Копия договора.\n2. Копия претензии.\n3. Расчёт суммы иска.')

    doc.add_paragraph()
    doc.add_paragraph('{{ lawyer_name }}  _______________  /  {{ data }} г.')

    path = OUT / 'iskovoe.docx'
    doc.save(path)
    print(f'✓ {path.name}')
    return path


if __name__ == '__main__':
    create_dogovor_uslug()
    create_dogovor_postavki()
    create_pretenziya()
    create_doverennost()
    create_iskovoe()
    print('\nВсе шаблоны созданы.')
